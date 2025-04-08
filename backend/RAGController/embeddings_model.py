"""
Módulo para la extracción de texto y creación de incrustaciones (embeddings) desde diferentes formatos de documentos.

Este módulo proporciona herramientas para extraer texto de archivos PDF, TXT y DOCX. Utiliza el modelo de
incrustación `sentence-transformers/multi-qa-MiniLM-L6-cos-v1` de Hugging Face para crear vectores
de texto a partir del contenido extraído de los documentos. Los documentos son divididos en fragmentos
pequeños para generar incrustaciones más manejables y relevantes.
"""

# Importación de Módulos
import fitz  # PyMuPDF para PDF
import docx
import io
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document
from langchain.chains import ConversationalRetrievalChain


# Declaración de Clases
class EmbeddingsModel:

    def __init__(self, model_name="sentence-transformers/multi-qa-MiniLM-L6-cos-v1", chunk_size=256, chunk_overlap=51):
        """
        Clase para la extracción de texto desde documentos y la creación de incrustaciones (embeddings) de texto.

        Esta clase facilita la extracción de texto desde archivos en formatos PDF, TXT y DOCX. También, esta clase
        puede crear incrustaciones utilizando el modelo de HuggingFace `sentence-transformers/multi-qa-MiniLM-L6-cos-v1`.
        """
        self.__embeddings_model_name = model_name
        # Se define el modelo de incrustaciones que se estará utilizando
        self._embeddings = HuggingFaceEmbeddings(model_name=self.__embeddings_model_name)
        # Se establece la cantidad de información que contendrá cada vector
        self._text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    @staticmethod
    def extract_pdf_text(file_content: bytes):
        """
        Método encargado de extraer el texto contenido en documentos PDF.

        Args:
            file_content (bytes): Contenido del documento.

        Returns:
            dict: un diccionario con el estado, mensaje y resultado de la operación.
                - 'status' (bool): Estado de la operación.
                - 'message' (str): Mensaje de éxito o error.
                - 'content' (str): Texto extraído del PDF.
        """

        text_buffer = io.StringIO()
        with fitz.open(stream=file_content, filetype="pdf") as doc:
            for page in doc:
                text_buffer.write(page.get_text())

        return {'status': True,
                'message': 'Se logró extraer el texto del documento PDF.',
                'content': text_buffer.getvalue()}

    @staticmethod
    def extract_txt_text(file_content: bytes):
        """
        Método encargado de extraer el texto contenido en documentos de texto.

        Args:
            file_content (bytes): Contenido del documento.

        Returns:
            dict: un diccionario con el estado, mensaje y resultado de la operación.
                - 'status' (bool): Estado de la operación.
                - 'message' (str): Mensaje de éxito o error.
                - 'content' (str): Texto extraído del documento de texto plano.
        """

        return {'status': True,
                'message': 'Se logró extraer el texto del documento de texto plano.',
                'content': file_content.decode("utf-8", errors="ignore")}

    @staticmethod
    def extract_docx_text(file_content: bytes):
        """
        Método encargado de extraer el texto contenido en documentos de Microsoft Word.

        Args:
            file_content (bytes): Contenido del documento.

        Returns:
            dict: un diccionario con el estado, mensaje y resultado de la operación.
                - 'status' (bool): Estado de la operación.
                - 'message' (str): Mensaje de éxito o error.
                - 'content' (str): Texto extraído del documento de Microsoft Word.
        """

        file_stream = io.BytesIO(file_content)
        doc = docx.Document(file_stream)

        return {'status': True,
                'message': 'Se logró extraer el texto del documento de Microsoft Word.',
                'content': "\n".join(para.text for para in doc.paragraphs if para.text.strip())}

    def process_document(self, file_content: bytes, file_extension: str):
        """
        Método encargado de extraer el texto de un documento dependiendo de su extensión de archivo.

        Args:
            file_content (bytes): Contenido el documento.
            file_extension (bytes): Extensión del documento.

        Returns:
            dict: un diccionario con el estado, mensaje y resultado de la operación.
                - 'status' (bool): Estado de la operación.
                - 'message' (str): Mensaje de éxito o error.
                - 'content' (str): Texto obtenido de la operación.
        """
        try:

            if file_extension == 'pdf':
                return self.extract_pdf_text(file_content)
            elif file_extension == 'txt':
                return self.extract_txt_text(file_content)
            elif file_extension == 'docx':
                return self.extract_docx_text(file_content)
            else:
                return {'status': False,
                        'message': 'El soporte a la extensión del documento que se ha ingresado no está disponible.'}

        except Exception as error:
            return {'status': False,
                    'message': f'Ocurrió un problema al extraer el texto del documento: f{str(error)}.'}

    def create_embedding(self, text: str, category: str):
        """
        Método encargado de crear las incrustaciones para un texto y asignarle los metadatos correspondientes.

        Args:
            text (str): El de texto a convertir en un vector.
            category (str): Los metadatos a asignar al documento que contendrá al vector.

        Returns:
            dict: un diccionario con el estado, mensaje y resultado de la operación.
                - 'status' (bool): Estado de la operación.
                - 'message' (str): Mensaje de éxito o error.
                - 'content' (str): La incrustación que resultó del texto proporcionado.
        """
        try:
            text_chunks = self._text_splitter.split_text(text)
            documents = [Document(page_content=chunk, metadata={"category": category}) for chunk in text_chunks]

            return {'status': True,
                    'message': 'Se creó la incrustación del texto de manera correcta.',
                    'content': documents}

        except Exception as error:
            return {'status': False,
                    'message': f'Ocurrió un problema al crear procesar el texto para la incrustación: {str(error)}.'}
