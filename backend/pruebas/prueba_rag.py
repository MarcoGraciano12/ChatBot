# IMPORTACIÓN DE MÓDULOS
import os
from dataclasses import fields

import fitz  # PyMuPDF para PDF
import docx
import mmap
import io
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document
import ollama


# Del colab
# imports de langchain, plotly y Chroma


# from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain

# from langchain_community.chat_models import ChatOllama
#
# from langchain_core.callbacks import StdOutCallbackHandler


class EmbeddingsModel:

    def __init__(self, model_name="sentence-transformers/multi-qa-MiniLM-L6-cos-v1", chunk_size=256, chunk_overlap=51):
        self.__embeddings_model_name = model_name
        # Se define el modelo de incrustaciones que se estará utilizando
        self._embeddings = HuggingFaceEmbeddings(model_name=self.__embeddings_model_name)
        # Se establece la cantidad de información que contendrá cada vector
        self._text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    @staticmethod
    def extract_pdf_text(file_content: bytes):
        """
        Extrae texto de un archivo PDF usando PyMuPDF.

        :return: Texto extraído del archivo PDF.
        """
        text_buffer = io.StringIO()
        with fitz.open(stream=file_content, filetype="pdf") as doc:
            for page in doc:
                text_buffer.write(page.get_text())
        return text_buffer.getvalue()

    @staticmethod
    def extract_txt_text(file_content: bytes):
        """
        Extrae texto de un archivo TXT.

        :return: Texto extraído del archivo TXT.
        """
        return file_content.decode("utf-8", errors="ignore")

    @staticmethod
    def extract_docx_text(file_content: bytes):
        """
        Extrae texto de un archivo DOCX.

        :return: Texto extraído del archivo DOCX.
        """
        file_stream = io.BytesIO(file_content)
        doc = docx.Document(file_stream)
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

    def process_document(self, file_content: bytes, file_extension: str):
        """
        Determina el tipo de archivo y llama al método adecuado para extraer el texto.

        :return: Texto extraído del archivo.
        :raises ValueError: Si el formato del archivo no es compatible.
        """
        try:
            if file_extension == 'pdf':
                return self.extract_pdf_text(file_content)
            elif file_extension == 'txt':
                return self.extract_txt_text(file_content)
            elif file_extension == 'docx':
                return self.extract_docx_text(file_content)
            else:
                print(f">>> Formato no soportado.")
                return None

        except Exception as e:
            print(f">>> Error al procesar el documento: {e}.")
            return None

    def create_embedding(self, text, category):
        try:
            text_chunks = self._text_splitter.split_text(text)
            documents = [Document(page_content=chunk, metadata={"category": category}) for chunk in text_chunks]
            return documents
        except Exception as e:
            print(f">>> Error al crear las incrustaciones: {e}.")
            return None



class ChromaDBManager(EmbeddingsModel):

    def __init__(self, db_name="cheeto", db_dir=r".\prueba"):
        super().__init__()
        self.__db_name = db_name    # Nombre de la base de datos
        self.__db_dir = db_dir  # Ubicación de la base de datos
        # Se asegura la existencia de la carpeta que contiene a la bd
        os.makedirs(self.__db_dir, exist_ok=True)
        # Se obtiene la ruta a la base de datos
        self.__db_path = os.path.join(self.__db_dir, self.__db_name)
        # Se inicializa la conexión a la base de datos
        self.__db = Chroma(persist_directory=self.__db_path, embedding_function=self._embeddings)
        self.__collections = self.load_categories()  # Lista con el nombre de las colecciones disponibles.

    def create_collection(self, file_content: bytes, file_extension: str, category):
        try:
            # Si el nombre de la colección ya existe
            if category in self.__collections:
                return False, f"El nombre del entrenamiento {category} ya está registrado."

            # Se procesa el documento.
            text = self.process_document(file_content, file_extension)
            print(text)
            if not text:    # Si no se obtuvo el texto del documento.
                return False, f"No se logró procesar el documento proporcionado."

            # Se obtienen pedazos del texto procesado con los metadatos asignados.
            documents = self.create_embedding(text=text, category=category)

            if not documents:   # Si no se obtuvieron los pedazos con metadatos.
                return False, f"No se lograron asignar los datos al entrenamiento correspondiente."

            # Se agrega el texto dividido como incrustaciones a la base de datos.
            self.__db.add_documents(documents)
            # Se agrega el nombre a las colecciones existentes.
            self.__collections.append(category)

            return True, f"Se ha creado el entrenamiento {category}."

        except Exception as error:
            print(f">>> Error al crear una colección: {error}.")
            return False, "Ocurrió un error inesperado al crear el entrenamiento."

    def delete_collection(self, category):
        try:
            if not category in self.__collections:  # Si la colección no existe.
                return False, f"El entrenamiento no está registrado."

            # Se eliminan los documentos con los metadatos correspondientes.
            self.__db.delete(where={"category": category})
            # Se elimina el nombre del colección de la lista de colecciones.
            self.__collections.remove(category)

            return True, f"Se eliminó el entrenamiento {category}."

        except Exception as error:
            print(f"Error al eliminar la colección: {error}.")
            return False, f"Ocurrió un error inesperado al borrar el entrenamiento."

    def update_collection(self, file_content: bytes, file_extension: str, category):
        try:
            if not category in self.__collections:  # Si la colección no existe.
                return False, f"El entrenamiento no está registrado."

            # Se procesa el documento.
            text = self.process_document(file_content, file_extension)

            if not text:  # Si no se obtuvo el texto del documento.
                return False, f"No se logró procesar el documento proporcionado."

            # Se obtienen pedazos del texto procesado con los metadatos asignados.
            documents = self.create_embedding(text=text, category=category)

            if not documents:  # Si no se obtuvieron los pedazos con metadatos.
                return False, f"No se lograron asignar los datos al entrenamiento correspondiente."

            # Se agrega el texto dividido como incrustaciones a la base de datos.
            self.__db.add_documents(documents)

            return True, f"Se ha actualizado el entrenamiento {category}."

        except Exception as error:
            print(f"Error al actualizar la colección: {error}.")
            return False, f"Ocurrió un error inesperado al actualizar el entrenamiento."

    def load_categories(self):
        try:
            # Se obtienen todos los documentos de la base de datos.
            all_docs = self.__db.get(include=['metadatas'])

            categories = set()

            # Se filtran los resultados con base a sus metadatos.
            for metadata in all_docs['metadatas']:
                if metadata and "category" in metadata:
                    categories.add(metadata["category"])

            return list(categories)

        except Exception as error:
            print(f"Error al obtener los entrenamientos: {error}.")
            return []

    def db_query(self, query, category, k):
        try:
            if not category in self.__collections:  # Si el entrenamiento no existe
                return False, f"El entrenamiento {category} no está registrado."

            # Se obtiene el resultado de la consulta a la base de datos.
            results = self.__db.similarity_search(query, k=k, filter={"category": category})

            if not results: # Si no se encontraron coincidencias.
                return False, "No se encontraron resultados para la consulta en la base de datos."

            return True, [res.page_content for res in results]

        except Exception as error:
            print(f">>> Error al consultar la base de datos: {error}.")
            return False, "Ocurrió un error inesperado al consultar la base de datos."

    def get_collections(self):
        """
        Método encargado de retornar la lista de colecciones disponibles en la base de datos.
        """
        return self.__collections



class OllamaSingleton:
    """
    Clase encargada de facilitar la implementación del patrón de diseño Singleton.

    La finalidad de esta clase es mantener una única instancia de ollama, reduciendo
    así tiempos de respuesta y costes de memoria del sistema.
    """
    _instance = None

    def __new__(cls, *args, **kwargs):
        """
        Crea una nueva instancia de la clase si no existe, de lo contrario, devuelve
        la instancia existente.
        """
        try:
            if cls._instance is None:
                cls._instance = super(OllamaSingleton, cls).__new__(cls, *args, **kwargs)
                cls._instance.client = ollama  # Mantener la instancia única de Ollama.
        except Exception as error:
            print(f">>> Error al manejar la instancia de Ollama: {error}.")
        finally:
            return cls._instance    # Se retorna la instancia de Ollama.



class ModelManager:
    """
    Clase encargada de gestionar y aplicar la configuración personalizada
    del usuario al modelo.
    """

    def __init__(self):
        # Se carga la lista de los modelos de Ollama disponibles.
        self.__available_models = self.load_models()

        # Variable para almacenar el modelo activo.
        self.__selected_model = None

        # Se obtiene la instancia de Ollama.
        self.ollama_instance = OllamaSingleton()

        # Se establece por defecto la cantidad de coincidencias del RAG a obtener.
        self._k = 1

        # Se establece por defecto la calidad de respuesta del modelo
        self._level = 0

        # Variable para gestionar la colección sobre la cual se realizarán las búsquedas.
        self._category = None

    @staticmethod
    def load_models():
        """
        Método destinado a obtener la lista de modelos de Ollama que se encuentran en
        local.
        """
        try:
            return [llm['model'] for llm in ollama.list()['models']]
        except Exception as error:
            print(f"Error al cargar la lista de modelos locales: {error}.")
            return []

    def get_list_models(self):
        """
        Método destinado a retornar la lista de modelos de Ollama que se encuentran en
        local
        """
        return self.__available_models

    def change_selected_model(self, index:int):
        try:
            if 0 <= index < len(self.__available_models):
                # Se asigna el nuevo modelo activo
                self.__selected_model = self.__available_models[index]
                return True, f"Se ha activado el modelo {self.__available_models[index]}."

            return False, f"El modelo seleccionado no se encuentra dentro de la lista de modelos disponibles."
        except Exception as error:
            print(f">>> Error al cambiar el modelo activo: {error}.")
            return False, "Ocurrió un error inesperado al intentar cambiar el modelo."

    def set_k(self, k:int):
        """
        Método para asignar un nuevo valor a la variable encargada de gestionar la
        cantidad de coincidencias del RAG.
        """
        self._k = k

    def set_level(self, level:int):
        """
        Método para asignar un nuevo valor a la variable encargada de gestionar la calidad
        de las respuestas del modelo.
        """
        self._level = level

    def set_category(self, category:str):
        """
        Método para asignar un nuevo valor a la variable encargada de gestionar
        la colección sobre la cual se realizarán las búsquedas en la base de datos.
        """
        self._category = category

    def get_selected_model(self):
        """
        Método para retornar al modelo activo.
        """
        return self.__selected_model

    def get_selected_category(self):
        """
        Método para retornar la categoría que se encuentra activa.
        """
        return self._category

    def get_k(self):
        """
        Método para retornar el número de coincidencias del RAG a obtener.
        """
        return self._k

    def get_level(self):
        """
        Método para retornar el nivel de respueta del modelo.
        """
        return self._level


class ChatSession:
    """Maneja la interacción con el modelo seleccionado."""

    def __init__(self):
        self._model_manager = ModelManager()
        self.db_manager = ChromaDBManager()

    # Obtener lista de modelos disponibles.
    def get_available_models(self):
        """
        Método que retornar la lista de modelos disponibles de Ollama.
        """
        return self._model_manager.get_list_models()

    # Obtener el modelo activo.
    def get_active_model(self):
        """
        Método que retorna al modelo que se encuentre activado.
        """
        model = self._model_manager.get_selected_model()

        if not model:
            return False, "No se ha seleccionado un modelo."

        return True, model

    # Cambiar Modelo
    def change_model(self, index:int):
        """
        Método que permite cambiar el modelo activo según el usuario indique.
        """
        return self._model_manager.change_selected_model(index)

    # Obtener lista de colecciones disponibles.
    def get_available_collections(self):
        """
        Método encargado de retornar la lista de colecciones disponibles en la base de datos.
        """
        return self.db_manager.get_collections()

    # Obtener la colección activa.
    def get_active_category(self):
        category = self._model_manager.get_selected_category()

        if not category:
            return False, "No se ha seleccionado un entrenamiento."

        return True, category

    # Cambiar la colección activa.
    def change_category(self, category:str):
        """
        Método para asignar una nueva colección activa.
        """
        self._model_manager.set_category(category=category)
        return True, f"Se ha seleccionado el entrenamiento {category}."

    # Obtener el número de coincidencias del RAG a obtener.
    def get_rag_number_matches(self):
        return True, self._model_manager.get_k()

    # Cambiar el número de coincidencias del RAG a obtener.
    def change_rag_number_matches(self, k:int):
        self._model_manager.set_k(k)
        return True, f"Se cambió el número de coincidencias del RAG a {k}."

    # Obtener el nivel de respuesta del modelo.
    def get_level_response(self):
        return True, self._model_manager.get_level()

    # Cambiar el nivel de respuesta del modelo.
    def change_level_response(self, level:int):
        self._model_manager.set_level(level)
        return True, f"Se cambió el nivel de respuesta del modelo."

    # Crear colección
    def create_collection(self, file_content: bytes, file_extension: str, category):
        """
        Método que permite la creación de una nueva colección (entrenamiento) en la base de datos.
        """
        return self.db_manager.create_collection(file_content=file_content,
                                                 file_extension=file_extension,
                                                 category=category)

    # Borrar Colección
    def delete_collection(self, category):
        """
        Método que permite eliminar una colección (entrenamiento) de la base de datos.
        """
        return self.db_manager.delete_collection(category=category)

    # Actualizar Colección
    def update_collection(self, file_content: bytes, file_extension: str, category):
        """
        Método que permite actualizar una colección (entrenamiento) de la base de datos.
        """
        return self.db_manager.update_collection(file_content=file_content,
                                                 file_extension=file_extension,
                                                 category=category)

    # Consultar Colección
    def query_collection(self, query, category):
        """
        Método para consultar una colección (entrenamiento) de la base de datos.
        """
        return self.db_manager.db_query(query=query, category=category, k=self._model_manager.get_k())

    # Validar configuración del chat
    def validate_model_settings(self):
        """
        Método encargado de validar si todas las opciones de configuración necesarias para interactuar con el chat
        están configuradas.
        """
        category = self._model_manager.get_selected_category()

        if not category:
            return False, "No se ha seleccionado un entrenamiento."

        model = self._model_manager.get_selected_model()

        if not model:
            return False, "No se ha seleccionado uno modelo de llm."

        return True, "Todo está listo para consultar el llm."


    # Chat con el modelo de Ollama activo en stream
    def query_ollama_model(self, query):
        """
        Método para consultar el modelo de ollama.
        """

        response = "\n\n".join(self.db_manager.db_query(query, self._model_manager.get_selected_category(),
                                            self._model_manager.get_k())[1])

        contexto = f"""
        
        Eres un asistente virtual muy inteligente y sofisticado.
        Las personas acuden a ti para satisfacer sus preguntas y dudas.
        Debes responser de manera directa y lo más breve posible, entre menos palabras tengas tus respuestas, mejor.
        Para responder a las preguntas, deberás tomar como contexto los resultados obtenidos a través de un
        sistema RAG, por lo que deberás analizar el contexto y responder de forma precisa.
        
        El contexto es:
        {response}
        """
        print(contexto)
        try:
            print(query)
            response = self._model_manager.ollama_instance.client.chat(
                model=self._model_manager.get_selected_model(),
                messages=[
                    {'role': 'system', 'content': contexto},
                    {'role': 'user', 'content': query},
                ],
                stream=True
            )
            for chunk in response:
                yield chunk['message']['content']
        except Exception as e:
            print(f"❌ Error al consultar el modelo: {e}")











    # def query_model(self, query: str, rag=1, level=1):
    #     """Envía la consulta al modelo y devuelve el stream de respuesta."""
    #     if not query.strip():
    #         print("⚠️ No puedes enviar un mensaje vacío.")
    #         return
    #
    #     response_level = ["debes acortar y simplificar tu respuesta lo más posible.",
    #                       "debes responder de forma directa y no tan extenso.",
    #                       "debes de proporcionar una respuesta expandida y sin inventar nada."]
    #
    #     db_consult = "\n".join(self.rag_handler.query_db(query,rag))
    #
    #     contexto = ("Eres un asistente virutal que pertenece a Grupo Fórmula. "
    #                 "El usuario te hará una pregunta y deberás responder solamente con base al contexto que se te"
    #                 "proporciona. El contexto que necesitas para responder la pregunta es"
    #                 f": {db_consult}. Recuerda: {response_level[level]}")
    #     try:
    #         response = self.ollama_instance.client.chat(
    #             model=self.model_manager.selected_model,
    #             messages=[
    #                 {'role': 'system', 'content': contexto},
    #                 {'role': 'user', 'content': query},
    #             ],
    #             stream=True
    #         )
    #         for chunk in response:
    #             yield chunk['message']['content']
    #     except Exception as e:
    #         print(f"❌ Error al consultar el modelo: {e}")

# Ejemplo de uso
if __name__ == "__main__":
    file_path = r"C:\Users\magraciano\Downloads\ELABORACION-DE-CHOCOLATE.pdf"
    manager = ChromaDBManager()
    # delete = manager.delete_collection(category="GF")
    # print(delete)
    # create = manager.create_collection(file_path=file_path, category="chocolate")
    # print(create)

    while True:
        user_input = input("Tu: ")

        if user_input.lower() == "salir":
            break

        result = manager.db_query(query=user_input, category="chocolate", k=10)[1]
        for response in result:
            print(f"-> {response}")
        print("====================================================================")


    # Modelo de embeddings (puedes cambiarlo por el que usas)
    # embedding_function = HuggingFaceEmbeddings(model_name="sentence-transformers/multi-qa-MiniLM-L6-cos-v1")
    #
    # # Definir diferentes colecciones para distintos temas
    # chroma_waffle = Chroma(persist_directory="./turin", collection_name="waffle",
    #                        embedding_function=embedding_function)
    # chroma_hotdog = Chroma(persist_directory="./turin", collection_name="hotdog",
    #                        embedding_function=embedding_function)
    #
    # # Agregar documentos a cada colección
    # chroma_waffle.add_texts(["Los waffles son deliciosos con miel."])
    # chroma_hotdog.add_texts(["Los hot dogs son populares en eventos deportivos."])
    #
    # # Consultar cada colección por separado
    # result_waffle = chroma_waffle.similarity_search("¿Qué se puede poner en un waffle?")
    # result_hotdog = chroma_hotdog.similarity_search("¿Dónde se venden más hot dogs?")
    #
    # print("Resultados para waffles:", result_waffle)
    # print("Resultados para hot dogs:", result_hotdog)