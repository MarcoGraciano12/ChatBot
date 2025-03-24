# IMPORTACIÓN DE MÓDULOS
import os
import fitz  # PyMuPDF para PDF
import docx
import mmap
import io
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document


class RAGManager:
    def __init__(self, base_dir=r".\database", model_name="sentence-transformers/multi-qa-MiniLM-L6-cos-v1"):
        self.base_dir = base_dir
        self.model_name = model_name
        self.embeddings = HuggingFaceEmbeddings(model_name=model_name)
        os.makedirs(base_dir, exist_ok=True)
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=150, chunk_overlap=0)
        self.db_path = os.path.join(self.base_dir, "db")
        try:
            self.db = Chroma(persist_directory=self.db_path, embedding_function=self.embeddings)
        except Exception as e:
            print(f"Error initializing Chroma: {e}")

    def extract_text(self, file_content: bytes, file_extension: str):
        """Extrae texto desde contenido en memoria basado en la extensión del archivo."""
        try:
            if file_extension == "pdf":
                text_buffer = io.StringIO()
                with fitz.open(stream=file_content, filetype="pdf") as doc:
                    for page in doc:
                        text_buffer.write(page.get_text())
                return text_buffer.getvalue()

            elif file_extension == "docx":
                file_stream = io.BytesIO(file_content)
                doc = docx.Document(file_stream)
                return "\n".join(para.text for para in doc.paragraphs if para.text.strip())

            elif file_extension == "txt":
                return file_content.decode("utf-8", errors="ignore")

            else:
                print(f">>> Formato no soportado: {file_extension}")
                return None

        except Exception as e:
            print(f"Error al extraer texto: {e}")
            return None

    def add_documents(self, category, file_content: bytes, file_extension: str):
        """Agrega un documento recibido en memoria a la base de datos."""
        text = self.extract_text(file_content, file_extension)
        if not text:
            print(f">>> ⚠️ No se pudo extraer texto del archivo recibido.")
            return False

        text_chunks = self.text_splitter.split_text(text)
        documents = [Document(page_content=chunk, metadata={"category": category})
                     for chunk in text_chunks]

        self.db.add_documents(documents)
        print(f"✅ Documento agregado a la categoría '{category}'.")
        return True

    def query_documents(self, query, category, top_k=5):
        """Consulta los documentos en la base de datos por categoría."""
        results = self.db.similarity_search(query, k=top_k, filter={"category": category})
        if not results:
            print(">>> No se encontraron resultados.")
            return False
        print(f">>> 📌 Resultados para '{category}':")
        # for res in results:
        #     print(f"- {res.page_content} (Fuente: {res.metadata.get('source', 'Desconocida')})")
        return [res.page_content for res in results]

    def delete_documents(self, category):
        """Elimina documentos de una categoría específica en la base de datos."""
        self.db.delete(where={"category": category})
        print(f">>> 🗑️ Documentos de '{category}' eliminados con éxito.")
        return True

    def get_categories(self):
        """Recupera todas las categorías únicas almacenadas en la base de datos."""
        try:
            print(self.db_path)
            all_docs = self.db.get(include=['metadatas'])
            categories = set()
            for metadata in all_docs['metadatas']:
                if metadata and "category" in metadata:
                    categories.add(metadata["category"])

            if not categories:
                print(">>> No hay categorías almacenadas.")
                return []
            print(">>> Categorías disponibles:")
            for category in categories:
                print(f"- {category}")
            return list(categories)
        except Exception as e:
            print(f">>> Error al obtener categorías: {e}")
            return []


if __name__ == "__main__":
    manager = RAGManager()
    menu = """
    1) Agregar documento a categoría.
    2) Consultar categoría.
    3) Eliminar categoría.
    4) Actualizar documentos de categoría.
    5) Mostrar categorías disponibles.
    0) Salir.
    """

    while True:
        try:
            print(menu)
            menu_option = int(input("Seleccione una opción del menú: "))

            match menu_option:
                case 0:
                    print("¡Que tenga un buen día!")
                    break
                case 1:
                    category = input("Ingrese la categoría: ")
                    file_path = input("Ingrese la ruta del archivo: ")
                    manager.add_documents(category, file_path)
                case 2:
                    category = input("Ingrese la categoría: ")
                    query = input("Ingrese su consulta: ")
                    manager.query_documents(query, category)
                case 3:
                    category = input("Ingrese la categoría a eliminar: ")
                    manager.delete_documents(category)
                case 4:
                    category = input("Ingrese la categoría: ")
                    file_path = input("Ingrese la ruta del archivo: ")
                    manager.update_documents(category, file_path)
                case 5:
                    manager.get_categories()
        except Exception as e:
            print(f">>> Error de tipo {e}.")
